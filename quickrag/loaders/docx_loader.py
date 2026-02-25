"""DOCX file loader for QuickRAG."""

from pathlib import Path

from quickrag.loaders.base import BaseLoader, LoadedDocument


class DocxLoader(BaseLoader):
    """Loader for Microsoft Word (.docx) files.

    Extracts text content from DOCX documents including paragraphs and tables.
    Requires the `python-docx` package.
    """

    EXTENSIONS = {".docx"}

    def supports(self, source: str | Path) -> bool:
        """Check if source is a DOCX file."""
        path = Path(source)
        return path.suffix.lower() in self.EXTENSIONS

    def load(self, source: str | Path) -> list[LoadedDocument]:
        """Load a DOCX file.

        Args:
            source: Path to the DOCX file.

        Returns:
            List with single LoadedDocument containing all text.
        """
        from docx import Document as DocxDocument

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        doc = DocxDocument(str(path))
        parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        if not parts:
            return []

        content = "\n\n".join(parts)

        return [
            LoadedDocument(
                content=content,
                source=str(path.absolute()),
                metadata={
                    "filename": path.name,
                    "extension": path.suffix,
                    "size_bytes": path.stat().st_size,
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                },
            )
        ]
